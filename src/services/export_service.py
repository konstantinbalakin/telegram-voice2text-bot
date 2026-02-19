"""Export service for generating transcription files in multiple formats."""

import io
import logging
import re
from typing import Any

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Pt  # type: ignore[import-untyped]

from src.exceptions import ExportError
from src.services.pdf_generator import PDFGenerator
from src.utils.markdown_utils import strip_markdown

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting transcription text to MD, TXT, PDF, DOCX."""

    def __init__(self, pdf_generator: PDFGenerator | None = None) -> None:
        self._pdf_generator = pdf_generator or PDFGenerator()

    def export(self, text: str, fmt: str, filename: str) -> io.BytesIO:
        """Dispatch export to the appropriate format handler.

        Args:
            text: Text content to export
            fmt: Export format (md, txt, pdf, docx)
            filename: Base filename without extension

        Returns:
            BytesIO object with .name set

        Raises:
            ValueError: If format is not supported
        """
        handlers = {
            "md": self.export_md,
            "txt": self.export_txt,
            "pdf": self.export_pdf,
            "docx": self.export_docx,
        }
        handler = handlers.get(fmt)
        if handler is None:
            raise ValueError(f"Unsupported export format: {fmt!r}")
        return handler(text, filename)

    def export_md(self, text: str, filename: str) -> io.BytesIO:
        """Export text as Markdown file (content preserved as-is).

        Args:
            text: Markdown text
            filename: Base filename without extension

        Returns:
            BytesIO with .name set to {filename}.md
        """
        buf = io.BytesIO(text.encode("utf-8"))
        buf.name = f"{filename}.md"
        return buf

    def export_txt(self, text: str, filename: str) -> io.BytesIO:
        """Export text as plain TXT with markdown stripped.

        Args:
            text: Markdown text
            filename: Base filename without extension

        Returns:
            BytesIO with .name set to {filename}.txt
        """
        plain = strip_markdown(text)
        buf = io.BytesIO(plain.encode("utf-8"))
        buf.name = f"{filename}.txt"
        return buf

    def export_pdf(self, text: str, filename: str) -> io.BytesIO:
        """Export text as PDF using the existing PDFGenerator.

        Args:
            text: Text content (may contain Markdown formatting)
            filename: Base filename without extension

        Returns:
            BytesIO with .name set to {filename}.pdf

        Raises:
            ExportError: If PDF generation fails
        """
        try:
            logger.info(f"Exporting PDF: {filename}")
            pdf_bytes = self._pdf_generator.generate_pdf(text)
            buf = io.BytesIO(pdf_bytes)
            buf.name = f"{filename}.pdf"
            logger.info(f"PDF export complete: {filename}.pdf ({len(pdf_bytes)} bytes)")
            return buf
        except Exception as e:
            logger.error(f"PDF export failed: {e}", exc_info=True)
            raise ExportError(
                f"PDF export failed: {e}", user_message="Не удалось создать PDF"
            ) from e

    def export_docx(self, text: str, filename: str) -> io.BytesIO:
        """Export text as DOCX with basic markdown formatting.

        Args:
            text: Markdown text
            filename: Base filename without extension

        Returns:
            BytesIO with .name set to {filename}.docx

        Raises:
            ExportError: If DOCX generation fails
        """
        try:
            logger.info(f"Exporting DOCX: {filename}")
            buf = self._markdown_to_docx(text, filename)
            logger.info(f"DOCX export complete: {filename}.docx")
            return buf
        except Exception as e:
            logger.error(f"DOCX export failed: {e}", exc_info=True)
            raise ExportError(
                f"DOCX export failed: {e}", user_message="Не удалось создать DOCX"
            ) from e

    def _markdown_to_docx(self, text: str, filename: str) -> io.BytesIO:
        """Convert markdown text to DOCX with basic formatting.

        Handles: H1-H3 headers, bullet lists, numbered lists, bold, italic, paragraphs.

        Args:
            text: Markdown text
            filename: Base filename without extension

        Returns:
            BytesIO with .name set to {filename}.docx
        """
        doc = Document()

        if not text:
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            buf.name = f"{filename}.docx"
            return buf

        lines = text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            if not stripped:
                i += 1
                continue

            header_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
            if header_match:
                level = len(header_match.group(1))
                header_text = header_match.group(2)
                doc.add_heading(self._strip_inline_markdown(header_text), level=level)
                i += 1
                continue

            if re.match(r"^[-•]\s+", stripped):
                while i < len(lines) and re.match(r"^[-•]\s+", lines[i].strip()):
                    item_text = re.sub(r"^[-•]\s+", "", lines[i].strip())
                    p = doc.add_paragraph(style="List Bullet")
                    self._add_formatted_text(p, item_text)
                    i += 1
                continue

            if re.match(r"^\d+\.\s+", stripped):
                while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
                    item_text = re.sub(r"^\d+\.\s+", "", lines[i].strip())
                    p = doc.add_paragraph(style="List Number")
                    self._add_formatted_text(p, item_text)
                    i += 1
                continue

            p = doc.add_paragraph()
            self._add_formatted_text(p, stripped)
            i += 1

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        buf.name = f"{filename}.docx"
        return buf

    def _strip_inline_markdown(self, text: str) -> str:
        """Strip inline markdown (bold, italic, code) from text."""
        result = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        result = re.sub(r"\*(.+?)\*", r"\1", result)
        result = re.sub(r"`(.+?)`", r"\1", result)
        return result

    def _add_formatted_text(self, paragraph: Any, text: str) -> None:
        """Add text with bold/italic/code formatting to a docx paragraph.

        Parses **bold**, *italic*, and `code` markers and applies formatting via runs.

        Args:
            paragraph: python-docx paragraph object
            text: Text that may contain **bold**, *italic*, and `code` markers
        """
        # Split by bold and italic markers
        # Pattern: **bold** or *italic*
        pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
        last_end = 0

        for match in pattern.finditer(text):
            # Add text before match
            if match.start() > last_end:
                run = paragraph.add_run(text[last_end : match.start()])
                run.font.size = Pt(11)

            if match.group(2):
                # Bold
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.font.size = Pt(11)
            elif match.group(3):
                # Italic
                run = paragraph.add_run(match.group(3))
                run.italic = True
                run.font.size = Pt(11)
            elif match.group(4):
                # Inline code
                run = paragraph.add_run(match.group(4))
                run.font.size = Pt(11)

            last_end = match.end()

        # Add remaining text
        if last_end < len(text):
            run = paragraph.add_run(text[last_end:])
            run.font.size = Pt(11)
